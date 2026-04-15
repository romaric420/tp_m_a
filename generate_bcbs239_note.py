"""
Generateur de la note reglementaire BCBS 239 — Format Word (.docx)
A convertir en PDF via Word ou un outil en ligne.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ============================================================
# EN-TETE
# ============================================================
header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header_para.add_run('QUANTAXIS CAPITAL')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub.add_run('Département Data Engineering — Division Risk Management')
run_sub.font.size = Pt(9)
run_sub.font.color.rgb = RGBColor(100, 100, 100)

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = line.add_run('─' * 60)
run_line.font.color.rgb = RGBColor(0, 51, 102)

# TITRE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title.add_run(
    'Note réglementaire BCBS 239\n'
    'Conformité du pipeline TradeCleanse aux principes\n'
    'd\'agrégation des données de risque'
)
run_t.bold = True
run_t.font.size = Pt(13)

# META
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run(
    f'Destinataire : Chief Risk Officer (CRO)\n'
    f'Classification : CONFIDENTIEL\n'
    f'Date : {datetime.date.today().strftime("%d/%m/%Y")}'
)
meta_run.font.size = Pt(9)
meta_run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ============================================================
# 1. CONTEXTE
# ============================================================
h1 = doc.add_heading('1. Contexte et objectif', level=2)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    'Le Comité de Bâle (BCBS) a publié en janvier 2013 le standard BCBS 239 '
    '« Principles for effective risk data aggregation and risk reporting », '
    'applicable aux G-SIBs et adopté comme bonne pratique par les régulateurs '
    'européens (BCE, ACPR). Ce cadre impose aux établissements financiers de '
    'garantir l\'exactitude, la complétude et la traçabilité des données '
    'alimentant les modèles de risque.'
)
doc.add_paragraph(
    'Le pipeline TradeCleanse a été développé pour auditer, nettoyer et '
    'certifier le dataset de 8 950 transactions consolidé depuis trois sources '
    '(Bloomberg, Murex, Refinitiv) avant son utilisation dans le modèle de '
    'scoring de risque de contrepartie. La présente note démontre la conformité '
    'de ce pipeline aux quatre principes BCBS 239 pertinents.'
)

# ============================================================
# 2. PRINCIPE 2 — Exactitude et intégrité
# ============================================================
h2 = doc.add_heading('2. Principe 2 — Exactitude et intégrité', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    '« Les données de risque agrégées doivent être exactes et fiables. '
    'Des contrôles doivent être en place pour garantir cette exactitude. »'
).runs[0].italic = True

doc.add_paragraph('Mesures implémentées dans TradeCleanse :')

table2 = doc.add_table(rows=6, cols=2, style='Light List Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
data2 = [
    ('Contrôle', 'Implémentation'),
    ('Valeurs sentinelles', '92 sentinelles détectées et remplacées (#N/A, #VALUE!, '
     '99999, nd) — Étape 1'),
    ('Doublons trade_id', '200 doublons supprimés (keep=first, booking Murex original) '
     '— Étape 2'),
    ('Cohérence bid/ask', '120 inversions corrigées par swap ; mid_price recalculé '
     'systématiquement — Étape 5'),
    ('Prix d\'exécution', '150 prix hors fourchette [bid×0.995, ask×1.005] remplacés '
     'par mid_price — Étape 5'),
    ('Rating vs défaut', '96 contradictions AAA/AA/A + default=1 corrigées par '
     'rétrogradation à BBB — Étape 5'),
]
for i, (c1, c2) in enumerate(data2):
    table2.rows[i].cells[0].text = c1
    table2.rows[i].cells[1].text = c2
    if i == 0:
        for cell in table2.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph(
    'Résultat : le Data Quality Score (DQS) passe de 98.6% (brut) à 100% '
    '(nettoyé). Les 14 expectations de validation sont toutes au statut PASS.'
)

# ============================================================
# 3. PRINCIPE 3 — Complétude
# ============================================================
h3 = doc.add_heading('3. Principe 3 — Complétude', level=2)
for run in h3.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    '« Les données de risque agrégées doivent couvrir l\'ensemble des '
    'expositions significatives. Les lacunes doivent être identifiées '
    'et documentées. »'
).runs[0].italic = True

doc.add_paragraph(
    'Le profiling initial a identifié 3 colonnes avec des valeurs manquantes '
    'significatives :'
)

table3 = doc.add_table(rows=4, cols=3, style='Light List Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
data3 = [
    ('Colonne', 'Taux NaN', 'Stratégie d\'imputation'),
    ('credit_rating', '14.98%', 'Mode par secteur GICS (plus pertinent que le mode '
     'global car le rating dépend du profil sectoriel)'),
    ('volatility_30d', '13.70%', 'Médiane (45.51%) + flag '
     'volatility_30d_was_missing'),
    ('country_risk', '0.34%', 'Médiane (38.10) + flag country_risk_was_missing'),
]
for i, (c1, c2, c3) in enumerate(data3):
    table3.rows[i].cells[0].text = c1
    table3.rows[i].cells[1].text = c2
    table3.rows[i].cells[2].text = c3
    if i == 0:
        for cell in table3.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph(
    'Chaque imputation est tracée via une colonne flag binaire (_was_missing) '
    'permettant au Risk Officer d\'identifier les données imputées dans les '
    'rapports de risque. La complétude finale atteint 100%.'
)

# ============================================================
# 4. PRINCIPE 6 — Adaptabilité
# ============================================================
h4 = doc.add_heading('4. Principe 6 — Adaptabilité', level=2)
for run in h4.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    '« Le processus d\'agrégation doit être capable de s\'adapter à des '
    'demandes ad hoc en période de stress. »'
).runs[0].italic = True

doc.add_paragraph(
    'Le pipeline TradeCleanse est conçu pour être réexécutable et adaptable :'
)

bullets_6 = [
    'Architecture modulaire : 10 étapes indépendantes avec logging structuré '
    '(horodatage, nb lignes avant/après). Chaque étape peut être désactivée '
    'ou modifiée sans impacter les autres.',
    'Monitoring du drift : le test de Kolmogorov-Smirnov (Bonus 2) compare '
    'automatiquement les distributions early/late et alerte en cas de '
    'changement de régime (p-value < 0.05).',
    'Détection d\'anomalies multivariées : l\'Isolation Forest flagge les '
    'combinaisons suspectes sans les supprimer, laissant la décision finale '
    'au Risk Officer.',
    'Pseudonymisation dynamique : le salt est lu depuis une variable '
    'd\'environnement (CLEANSE_SALT), permettant la rotation des clés '
    'sans modification du code source (conformité RGPD Art. 25).',
]
for b in bullets_6:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 5. PRINCIPE 8 — Exactitude du reporting
# ============================================================
h5 = doc.add_heading('5. Principe 8 — Exactitude du reporting', level=2)
for run in h5.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    '« Les rapports de risque doivent refléter fidèlement les données '
    'agrégées et être soumis à des procédures de validation. »'
).runs[0].italic = True

doc.add_paragraph(
    'Le pipeline produit une chaîne de validation complète :'
)

bullets_8 = [
    'Suite de 14 expectations automatisées (03_validation.py) couvrant '
    'l\'unicité, la cohérence, les plages de validité et l\'absence de PII. '
    'Score obtenu : 14/14.',
    'Rapport de qualité (Étape 10) avec métriques avant/après : lignes, '
    'colonnes, taux de complétude, doublons résiduels, DQS.',
    'Fichier de log horodaté (tradecleanse_pipeline.log) traçant chaque '
    'décision de nettoyage avec sa justification métier — piste d\'audit '
    'complète pour les régulateurs.',
    'Validation de l\'impact ML (Bonus 3) : comparaison Random Forest '
    'raw vs clean démontrant que le nettoyage améliore l\'AUC-ROC de +1.5% '
    'sur la prédiction de défaut.',
]
for b in bullets_8:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# CONCLUSION
# ============================================================
h6 = doc.add_heading('6. Conclusion et recommandations', level=2)
for run in h6.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    'Le pipeline TradeCleanse satisfait les exigences des principes 2, 3, 6 '
    'et 8 du standard BCBS 239. Le dataset nettoyé (8 750 lignes, DQS = 100%) '
    'est certifié pour alimenter le modèle de scoring de risque de contrepartie.'
)

doc.add_paragraph('Recommandations pour les prochaines itérations :')
reco = [
    'Intégrer Great Expectations en production pour automatiser les validations '
    'à chaque ingestion de données.',
    'Mettre en place un monitoring de drift en temps réel avec alertes '
    'automatiques vers le desk Risk.',
    'Étendre la détection de wash trading aux paires cross-trader (collusion).',
    'Planifier un audit externe annuel de la qualité des données conformément '
    'aux recommandations BCE/ACPR.',
]
for r in reco:
    doc.add_paragraph(r, style='List Bullet')

# SIGNATURE
doc.add_paragraph()
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig_run = sig.add_run(
    'Data Engineering Team — QuantAxis Capital\n'
    f'{datetime.date.today().strftime("%d/%m/%Y")}'
)
sig_run.font.size = Pt(9)
sig_run.italic = True

# SAUVEGARDE
output_path = 'BCBS239_Note_Reglementaire.docx'
doc.save(output_path)
print(f"Document Word genere : {output_path}")
print("Ouvrez-le dans Word et exportez en PDF (Fichier > Enregistrer sous > PDF)")
